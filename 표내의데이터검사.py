import os
from docx import Document

def search_word_files(folder_path, search_data):
    file_count = 0
    file_names = []

    for entry in os.scandir(folder_path):
        if entry.is_file() and entry.name.endswith(".docx"):
            try:
                doc = Document(entry.path)
                data_found = False

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if search_data.lower() in paragraph.text.lower():
                                    data_found = True
                                    break

                if data_found:
                    file_names.append(entry.name)
                    
            except Exception as e:
                print(f"오류 발생: {entry.name} - {str(e)}")
    
    file_count = len(file_names)
    return file_count, file_names

# 사용자로부터 검색 데이터 입력 받기
search_data = input("찾고자 하는 데이터를 입력하세요: ")

# 폴더 경로 설정
folder_path = "/Users/parkdahye/Desktop/ul"

import time
start_time = time.time()

# 검사 수행
file_count, file_names = search_word_files(folder_path, search_data)

end_time = time.time()
execution_time = round((end_time - start_time) / 60, 2)

# 결과 출력
total_file_count = sum(1 for _ in os.scandir(folder_path) if _.is_file())
print(f"검사한 대상 파일은 {total_file_count}개이며, 해당 데이터가 있는 파일의 갯수는 {file_count}개입니다.")
print(f"해당 데이터가 있는 파일의 파일명은 다음과 같습니다:")
for file_name in file_names:
    print(file_name)
print(f"해당 작업을 수행하는데 걸린 시간은 {execution_time}분 입니다.")
